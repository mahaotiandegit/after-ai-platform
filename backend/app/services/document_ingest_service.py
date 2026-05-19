from __future__ import annotations

import json
import re
import uuid
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader
from sqlalchemy import text
from sqlalchemy.orm import Session


UPLOAD_ROOT = Path("uploads/documents")

SUPPORTED_EXTENSIONS = {
    ".txt",
    ".md",
    ".markdown",
    ".csv",
    ".json",
    ".log",
    ".pdf",
    ".docx",
    ".xlsx",
}


class UnsupportedDocumentTypeError(ValueError):
    pass


def _safe_filename(filename: str | None) -> str:
    raw = Path(filename or "uploaded.txt").name
    cleaned = re.sub(r"[^A-Za-z0-9._\-\u4e00-\u9fff]", "_", raw).strip("._")
    return cleaned[:180] or "uploaded.txt"


def _file_type(filename: str) -> str:
    suffix = Path(filename).suffix.lower().lstrip(".")
    return (suffix or "txt")[:32]


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("文件编码无法识别，请先转成 UTF-8 文本后再上传。")


def _normalize_text(content: str) -> str:
    content = content.replace("\r\n", "\n").replace("\r", "\n")
    content = re.sub(r"\n{3,}", "\n\n", content)
    return content.strip()


def _extract_plain_text(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    content = _decode_text(data)

    if suffix == ".json":
        try:
            obj = json.loads(content)
            content = json.dumps(obj, ensure_ascii=False, indent=2)
        except Exception:
            pass

    return _normalize_text(content)


def _extract_pdf_text(data: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(data))
    except Exception as exc:
        raise ValueError(f"PDF 文件读取失败：{exc}") from exc

    if getattr(reader, "is_encrypted", False):
        raise ValueError("暂不支持加密 PDF，请先解除密码后上传。")

    pages: list[str] = []

    for index, page in enumerate(reader.pages, start=1):
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""

        page_text = _normalize_text(page_text)

        if page_text:
            pages.append(f"[第 {index} 页]\n{page_text}")

    content = "\n\n".join(pages).strip()

    if not content:
        raise ValueError("PDF 没有解析出文本。可能是扫描版 PDF，后续需要接 OCR。")

    return content


def _extract_docx_text(data: bytes) -> str:
    try:
        doc = DocxDocument(BytesIO(data))
    except Exception as exc:
        raise ValueError(f"Word docx 文件读取失败：{exc}") from exc

    parts: list[str] = []

    for paragraph in doc.paragraphs:
        text_value = paragraph.text.strip()
        if text_value:
            parts.append(text_value)

    for table_index, table in enumerate(doc.tables, start=1):
        rows: list[str] = []

        for row in table.rows:
            cells = [
                cell.text.strip().replace("\n", " ")
                for cell in row.cells
            ]

            while cells and not cells[-1]:
                cells.pop()

            if any(cells):
                rows.append(" | ".join(cells))

        if rows:
            parts.append(f"[表格 {table_index}]\n" + "\n".join(rows))

    content = "\n\n".join(parts).strip()

    if not content:
        raise ValueError("Word docx 没有解析出有效文本。")

    return _normalize_text(content)


def _extract_xlsx_text(data: bytes) -> str:
    try:
        workbook = load_workbook(
            filename=BytesIO(data),
            data_only=True,
            read_only=True,
        )
    except Exception as exc:
        raise ValueError(f"Excel xlsx 文件读取失败：{exc}") from exc

    sheets: list[str] = []

    for sheet in workbook.worksheets:
        rows: list[str] = []

        for row_index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            values: list[str] = []

            for value in row:
                if value is None:
                    values.append("")
                else:
                    values.append(str(value).strip())

            while values and not values[-1]:
                values.pop()

            if any(values):
                rows.append(f"R{row_index}: " + " | ".join(values))

        if rows:
            sheets.append(f"[工作表：{sheet.title}]\n" + "\n".join(rows))

    content = "\n\n".join(sheets).strip()

    if not content:
        raise ValueError("Excel xlsx 没有解析出有效文本。")

    return _normalize_text(content)


def _extract_text(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()

    if suffix not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise UnsupportedDocumentTypeError(
            f"当前支持文件类型：{supported}。旧版 .doc / .xls 暂不支持，请先另存为 .docx / .xlsx。"
        )

    if suffix in {".txt", ".md", ".markdown", ".csv", ".json", ".log"}:
        return _extract_plain_text(filename, data)

    if suffix == ".pdf":
        return _extract_pdf_text(data)

    if suffix == ".docx":
        return _extract_docx_text(data)

    if suffix == ".xlsx":
        return _extract_xlsx_text(data)

    raise UnsupportedDocumentTypeError(f"暂不支持的文件类型：{suffix}")


def _split_text(content: str, max_chars: int = 900, overlap: int = 120) -> list[str]:
    clean = content.strip()
    if not clean:
        return []

    if len(clean) <= max_chars:
        return [clean]

    chunks: list[str] = []
    start = 0

    while start < len(clean):
        end = min(start + max_chars, len(clean))

        if end < len(clean):
            split_candidates = [
                clean.rfind("\n\n", start, end),
                clean.rfind("\n", start, end),
                clean.rfind("。", start, end),
                clean.rfind(".", start, end),
                clean.rfind("；", start, end),
                clean.rfind(";", start, end),
            ]
            split_at = max(split_candidates)
            if split_at > start + max_chars * 0.5:
                end = split_at + 1

        chunk = clean[start:end].strip()
        if chunk:
            chunks.append(chunk)

        if end >= len(clean):
            break

        start = max(end - overlap, start + 1)

    return chunks


def ingest_document_bytes(
    db: Session,
    *,
    filename: str,
    content: bytes,
    title: str | None = None,
    content_type: str | None = None,
) -> dict:
    if not content:
        raise ValueError("上传文件为空。")

    safe_filename = _safe_filename(filename)
    file_type = _file_type(safe_filename)
    extracted_text = _extract_text(safe_filename, content)
    chunks = _split_text(extracted_text)

    if not chunks:
        raise ValueError("文件没有解析出有效文本。")

    document_id = str(uuid.uuid4())
    document_title = (title or Path(safe_filename).stem).strip()[:255] or safe_filename

    storage_dir = UPLOAD_ROOT / document_id
    storage_dir.mkdir(parents=True, exist_ok=True)

    storage_path = storage_dir / safe_filename
    storage_path.write_bytes(content)

    try:
        db.execute(
            text(
                """
                INSERT INTO documents (
                    id,
                    title,
                    file_name,
                    file_type,
                    storage_path,
                    status,
                    uploaded_by_id,
                    created_at,
                    updated_at
                )
                VALUES (
                    :id,
                    :title,
                    :file_name,
                    :file_type,
                    :storage_path,
                    'processing',
                    NULL,
                    now(),
                    now()
                )
                """
            ),
            {
                "id": document_id,
                "title": document_title,
                "file_name": safe_filename,
                "file_type": file_type,
                "storage_path": str(storage_path),
            },
        )

        for index, chunk in enumerate(chunks):
            chunk_metadata = {
                "source": "uploaded",
                "section": document_title,
                "policy_code": f"UPLOAD-{index + 1:03d}",
                "original_file_name": safe_filename,
                "content_type": content_type,
                "parser": file_type,
            }

            db.execute(
                text(
                    """
                    INSERT INTO document_chunks (
                        id,
                        document_id,
                        chunk_index,
                        content,
                        page_no,
                        token_count,
                        metadata,
                        created_at
                    )
                    VALUES (
                        :id,
                        :document_id,
                        :chunk_index,
                        :content,
                        NULL,
                        :token_count,
                        CAST(:metadata_json AS jsonb),
                        now()
                    )
                    """
                ),
                {
                    "id": str(uuid.uuid4()),
                    "document_id": document_id,
                    "chunk_index": index,
                    "content": chunk,
                    "token_count": max(1, len(chunk) // 2),
                    "metadata_json": json.dumps(chunk_metadata, ensure_ascii=False),
                },
            )

        db.execute(
            text(
                """
                UPDATE documents
                SET status = 'indexed',
                    updated_at = now()
                WHERE id = :id
                """
            ),
            {"id": document_id},
        )

        db.commit()

    except Exception:
        db.rollback()
        raise

    return {
        "document_id": document_id,
        "title": document_title,
        "file_name": safe_filename,
        "file_type": file_type,
        "storage_path": str(storage_path),
        "status": "indexed",
        "chunk_count": len(chunks),
        "total_chars": len(extracted_text),
    }


def _jsonb_param(value:dict)->str:
    return json.dumps(value or {},ensure_ascii=False)

def _to_iso(value)->str |None:
    if value is None:
        return None
    if hasattr(value,"isoformat"):
        return value.isoformat()
    
    return str(value)

def _task_row_to_dict(row:dict)->dict:
    result_metadata=row.get("result_metadata") or {}

    if isinstance(result_metadata,str):
        try:
            result_metadata=json.loads(result_metadata)
        except Exception:
            result_metadata={"raw":result_metadata}
    
    return {
        "task_id": str(row.get("task_id")),
        "document_id": str(row.get("document_id")),
        "title": row.get("title"),
        "file_name": row.get("file_name"),
        "file_type": row.get("file_type"),
        "storage_path": row.get("storage_path"),
        "task_type": row.get("task_type"),
        "task_status": row.get("task_status"),
        "document_status": row.get("document_status"),
        "attempt_count": int(row.get("attempt_count") or 0),
        "max_attempts": int(row.get("max_attempts") or 0),
        "error_message": row.get("error_message"),
        "result_metadata": result_metadata,
        "chunk_count": int(row.get("chunk_count") or 0),
        "total_chars": result_metadata.get("total_chars"),
        "created_at": _to_iso(row.get("created_at")),
        "started_at": _to_iso(row.get("started_at")),
        "finished_at": _to_iso(row.get("finished_at")),
        "updated_at": _to_iso(row.get("updated_at")),
    }

def create_document_index_task(
    db: Session,
    *,
    filename: str,
    content: bytes,
    title: str | None = None,
    content_type: str | None = None,
)->dict:
    if not content:
        raise ValueError("上传文件为空。")
    
    safe_filename=_safe_filename(filename)
    suffix=Path(safe_filename).suffix.lower()

    if suffix not in SUPPORTED_EXTENSIONS:
        supported=",".join(sorted(SUPPORTED_EXTENSIONS))
        raise UnsupportedDocumentTypeError(
             f"当前支持文件类型：{supported}。旧版 .doc / .xls 暂不支持，请先另存为 .docx / .xlsx。"
        )
    
    file_type=_file_type(safe_filename)
    document_id=str(uuid.uuid4())
    task_id = str(uuid.uuid4())
    document_title = (title or Path(safe_filename).stem).strip()[:255] or safe_filename

    storage_dir = UPLOAD_ROOT / document_id
    storage_dir.mkdir(parents=True, exist_ok=True)

    storage_path = storage_dir / safe_filename
    storage_path.write_bytes(content)

    task_metadata = {
        "source": "uploaded",
        "original_file_name": safe_filename,
        "content_type": content_type,
        "parser": file_type,
    }
    try:
        db.execute(
            text(
                """
                INSERT INTO documents (
                    id,
                    title,
                    file_name,
                    file_type,
                    storage_path,
                    status,
                    uploaded_by_id,
                    created_at,
                    updated_at
                )
                VALUES (
                    :id,
                    :title,
                    :file_name,
                    :file_type,
                    :storage_path,
                    'uploaded',
                    NULL,
                    now(),
                    now()
                )
                """
            ),
            {
                "id": document_id,
                "title": document_title,
                "file_name": safe_filename,
                "file_type": file_type,
                "storage_path": str(storage_path),
            },
        )

        db.execute(
            text(
                """
                INSERT INTO document_index_tasks (
                    id,
                    document_id,
                    task_type,
                    status,
                    attempt_count,
                    max_attempts,
                    error_message,
                    result_metadata,
                    created_at,
                    updated_at
                )
                VALUES (
                    :id,
                    :document_id,
                    'index_document',
                    'pending',
                    0,
                    3,
                    NULL,
                    CAST(:result_metadata AS jsonb),
                    now(),
                    now()
                )
                """
            ),
            {
                "id": task_id,
                "document_id": document_id,
                "result_metadata": _jsonb_param(task_metadata),
            },
        )

        db.commit()

    except Exception:
        db.rollback()
        raise

    return {
        "document_id": document_id,
        "task_id": task_id,
        "title": document_title,
        "file_name": safe_filename,
        "file_type": file_type,
        "storage_path": str(storage_path),
        "status": "pending",
        "document_status": "uploaded",
        "task_status": "pending",
        "chunk_count": 0,
        "total_chars": 0,
    }

def _read_task_with_document(db:Session,task_id:str)->dict:
    row=db.execute(
        text(
            """
            SELECT
                t.id AS task_id,
                t.document_id AS document_id,
                t.task_type AS task_type,
                t.status AS task_status,
                t.attempt_count AS attempt_count,
                t.max_attempts AS max_attempts,
                t.error_message AS error_message,
                t.result_metadata AS result_metadata,
                t.created_at AS created_at,
                t.started_at AS started_at,
                t.finished_at AS finished_at,
                t.updated_at AS updated_at,
                d.title AS title,
                d.file_name AS file_name,
                d.file_type AS file_type,
                d.storage_path AS storage_path,
                d.status AS document_status
            FROM document_index_tasks t
            JOIN documents d ON d.id = t.document_id
            WHERE t.id = :task_id
            """
        ),
        {"task_id": task_id},
    ).mappings().first()
    if row is None:
        raise ValueError("文档索引任务不存在。")

    return dict(row)

def get_document_index_task(db: Session, task_id: str) -> dict:
    row = db.execute(
        text(
            """
            SELECT
                t.id AS task_id,
                t.document_id AS document_id,
                t.task_type AS task_type,
                t.status AS task_status,
                t.attempt_count AS attempt_count,
                t.max_attempts AS max_attempts,
                t.error_message AS error_message,
                t.result_metadata AS result_metadata,
                t.created_at AS created_at,
                t.started_at AS started_at,
                t.finished_at AS finished_at,
                t.updated_at AS updated_at,
                d.title AS title,
                d.file_name AS file_name,
                d.file_type AS file_type,
                d.storage_path AS storage_path,
                d.status AS document_status,
                COUNT(c.id) AS chunk_count
            FROM document_index_tasks t
            JOIN documents d ON d.id = t.document_id
            LEFT JOIN document_chunks c ON c.document_id = d.id
            WHERE t.id = :task_id
            GROUP BY
                t.id,
                d.id
            """
        ),
        {"task_id": task_id},
    ).mappings().first()

    if row is None:
        raise ValueError("文档索引任务不存在。")

    return _task_row_to_dict(dict(row))

def list_document_index_tasks(
    db: Session,
    *,
    limit: int = 20,
    offset: int = 0,
) -> list[dict]:
    rows = db.execute(
        text(
            """
            SELECT
                t.id AS task_id,
                t.document_id AS document_id,
                t.task_type AS task_type,
                t.status AS task_status,
                t.attempt_count AS attempt_count,
                t.max_attempts AS max_attempts,
                t.error_message AS error_message,
                t.result_metadata AS result_metadata,
                t.created_at AS created_at,
                t.started_at AS started_at,
                t.finished_at AS finished_at,
                t.updated_at AS updated_at,
                d.title AS title,
                d.file_name AS file_name,
                d.file_type AS file_type,
                d.storage_path AS storage_path,
                d.status AS document_status,
                COUNT(c.id) AS chunk_count
            FROM document_index_tasks t
            JOIN documents d ON d.id = t.document_id
            LEFT JOIN document_chunks c ON c.document_id = d.id
            GROUP BY
                t.id,
                d.id
            ORDER BY t.created_at DESC
            LIMIT :limit OFFSET :offset
            """
        ),
        {"limit": limit, "offset": offset},
    ).mappings().all()

    return [_task_row_to_dict(dict(row)) for row in rows]

def run_document_index_task(db: Session, task_id: str) -> dict:
    task = _read_task_with_document(db, task_id)

    task_status = task.get("task_status")
    attempt_count = int(task.get("attempt_count") or 0)
    max_attempts = int(task.get("max_attempts") or 3)

    if task_status == "running":
        raise ValueError("文档索引任务正在执行中。")

    if task_status == "succeeded":
        return get_document_index_task(db, task_id)

    if attempt_count >= max_attempts:
        raise ValueError("文档索引任务已达到最大重试次数。")

    document_id = str(task["document_id"])
    storage_path = Path(str(task["storage_path"]))
    file_name = str(task["file_name"])
    file_type = str(task["file_type"])
    document_title = str(task["title"])

    try:
        db.execute(
            text(
                """
                UPDATE document_index_tasks
                SET status = 'running',
                    attempt_count = attempt_count + 1,
                    error_message = NULL,
                    started_at = now(),
                    finished_at = NULL,
                    updated_at = now()
                WHERE id = :task_id
                """
            ),
            {"task_id": task_id},
        )

        db.execute(
            text(
                """
                UPDATE documents
                SET status = 'indexing',
                    updated_at = now()
                WHERE id = :document_id
                """
            ),
            {"document_id": document_id},
        )

        db.commit()

        if not storage_path.exists():
            raise ValueError(f"文档源文件不存在：{storage_path}")

        content = storage_path.read_bytes()
        extracted_text = _extract_text(file_name, content)
        chunks = _split_text(extracted_text)

        if not chunks:
            raise ValueError("文件没有解析出有效文本。")

        db.execute(
            text("DELETE FROM document_chunks WHERE document_id = :document_id"),
            {"document_id": document_id},
        )

        for index, chunk in enumerate(chunks):
            chunk_metadata = {
                "source": "document_index_task",
                "task_id": task_id,
                "section": document_title,
                "policy_code": f"UPLOAD-{index + 1:03d}",
                "original_file_name": file_name,
                "parser": file_type,
            }

            db.execute(
                text(
                    """
                    INSERT INTO document_chunks (
                        id,
                        document_id,
                        chunk_index,
                        content,
                        page_no,
                        token_count,
                        metadata,
                        created_at
                    )
                    VALUES (
                        :id,
                        :document_id,
                        :chunk_index,
                        :content,
                        NULL,
                        :token_count,
                        CAST(:metadata_json AS jsonb),
                        now()
                    )
                    """
                ),
                {
                    "id": str(uuid.uuid4()),
                    "document_id": document_id,
                    "chunk_index": index,
                    "content": chunk,
                    "token_count": max(1, len(chunk) // 2),
                    "metadata_json": _jsonb_param(chunk_metadata),
                },
            )

        result_metadata = {
            "chunk_count": len(chunks),
            "total_chars": len(extracted_text),
            "parser": file_type,
            "storage_path": str(storage_path),
        }

        db.execute(
            text(
                """
                UPDATE documents
                SET status = 'indexed',
                    updated_at = now()
                WHERE id = :document_id
                """
            ),
            {"document_id": document_id},
        )

        db.execute(
            text(
                """
                UPDATE document_index_tasks
                SET status = 'succeeded',
                    error_message = NULL,
                    result_metadata = CAST(:result_metadata AS jsonb),
                    finished_at = now(),
                    updated_at = now()
                WHERE id = :task_id
                """
            ),
            {
                "task_id": task_id,
                "result_metadata": _jsonb_param(result_metadata),
            },
        )

        db.commit()

        return get_document_index_task(db, task_id)

    except Exception as exc:
        db.rollback()

        error_message = str(exc)[:2000]

        try:
            db.execute(
                text(
                    """
                    UPDATE documents
                    SET status = 'index_failed',
                        updated_at = now()
                    WHERE id = :document_id
                    """
                ),
                {"document_id": document_id},
            )

            db.execute(
                text(
                    """
                    UPDATE document_index_tasks
                    SET status = 'failed',
                        error_message = :error_message,
                        finished_at = now(),
                        updated_at = now()
                    WHERE id = :task_id
                    """
                ),
                {
                    "task_id": task_id,
                    "error_message": error_message,
                },
            )

            db.commit()

        except Exception:
            db.rollback()

        raise


def retry_document_index_task(db: Session, task_id: str) -> dict:
    task = _read_task_with_document(db, task_id)

    if task.get("task_status") == "running":
        raise ValueError("文档索引任务正在执行中，不能重试。")

    db.execute(
        text(
            """
            UPDATE document_index_tasks
            SET status = 'pending',
                attempt_count = 0,
                error_message = NULL,
                result_metadata = CAST(:result_metadata AS jsonb),
                started_at = NULL,
                finished_at = NULL,
                updated_at = now()
            WHERE id = :task_id
            """
        ),
        {
            "task_id": task_id,
            "result_metadata": _jsonb_param({"retry_reset": True}),
        },
    )

    db.execute(
        text(
            """
            UPDATE documents
            SET status = 'uploaded',
                updated_at = now()
            WHERE id = :document_id
            """
        ),
        {"document_id": str(task["document_id"])},
    )

    db.commit()

    return run_document_index_task(db, task_id)

def run_pending_document_index_tasks(db: Session, *, limit: int = 5) -> dict:
    rows = db.execute(
        text(
            """
            SELECT id
            FROM document_index_tasks
            WHERE status = 'pending'
               OR (status = 'failed' AND attempt_count < max_attempts)
            ORDER BY created_at ASC
            LIMIT :limit
            """
        ),
        {"limit": limit},
    ).scalars().all()

    items: list[dict] = []

    for item_task_id in rows:
        item_task_id = str(item_task_id)

        try:
            task_result = run_document_index_task(db, item_task_id)
            items.append(
                {
                    "task_id": item_task_id,
                    "ok": True,
                    "task": task_result,
                }
            )

        except Exception as exc:
            items.append(
                {
                    "task_id": item_task_id,
                    "ok": False,
                    "error": str(exc),
                }
            )

    return {
        "total": len(items),
        "items": items,
    }